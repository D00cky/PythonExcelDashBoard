from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from openpyxl.workbook import Workbook

from app.core.templates.sabesp_pimentas import SabespPimentasTemplate


def render_docx(template: SabespPimentasTemplate, workbook: Workbook, path: Path) -> bytes:
    inspections = template.extract_inspections(path)
    iqs_rows = template.extract_iqs_by_service(workbook)
    ic_rows = template.extract_ic_by_service(workbook)
    iqs_overall = template.extract_iqs_overall(workbook)
    periodo = _periodo(inspections) or template.extract_periodo(workbook)

    doc = Document()
    title = doc.add_heading(f"Dashboard SABESP — Polo {template.polo_name.title()}", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x26, 0x46, 0x53)

    if periodo:
        doc.add_paragraph().add_run(f"Período: {periodo}").bold = True
    if iqs_overall is not None:
        doc.add_paragraph().add_run(f"IQS Geral: {iqs_overall:.1%}").bold = True
    if not inspections.empty:
        doc.add_paragraph(f"Total de inspeções: {len(inspections)}")
        doc.add_paragraph(f"Equipes distintas: {inspections['team'].nunique()}")

    if ic_rows:
        doc.add_heading("Índice de Conformidade por Serviço", level=1)
        _table(
            doc,
            ["Serviço", "IC (%)", "LVs"],
            [[r.name, f"{r.ic_pct:.1%}", str(r.lvs)] for r in ic_rows],
        )

    if iqs_rows:
        doc.add_heading("Índice de Qualidade SABESP por Serviço", level=1)
        _table(
            doc,
            ["Serviço", "Avaliadas", "NC", "Conforme", "NC (%)", "Conforme (%)"],
            [
                [
                    r.name,
                    str(r.fotos_avaliadas),
                    str(r.fotos_nc),
                    str(r.fotos_conforme),
                    f"{r.nc_pct:.1%}",
                    f"{r.conforme_pct:.1%}",
                ]
                for r in iqs_rows
            ],
        )

    if not inspections.empty:
        for service in sorted(template.SERVICE_SHEETS):
            sub = inspections[inspections["service"] == service]
            if sub.empty:
                continue
            doc.add_heading(f"{service} — Top 10 Equipes", level=1)
            agg = sub.groupby("team").agg(
                conforme=("conforme_count", "sum"),
                nao_conforme=("nao_conforme_count", "sum"),
            )
            agg["total"] = agg["conforme"] + agg["nao_conforme"]
            top = agg[agg["total"] > 0].sort_values("total", ascending=False).head(10)
            _table(
                doc,
                ["Equipe", "Conforme", "Não Conforme"],
                [[team, str(int(r["conforme"])), str(int(r["nao_conforme"]))] for team, r in top.iterrows()],
            )

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _table(doc, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for c, h in enumerate(headers):
        cell = t.rows[0].cells[c]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            t.rows[r].cells[c].text = val


def _periodo(df) -> str | None:
    if df.empty or "start_date" not in df.columns:
        return None
    dates = df["start_date"].dropna()
    if dates.empty:
        return None
    return f"{dates.min():%d/%m/%Y} à {dates.max():%d/%m/%Y}"
